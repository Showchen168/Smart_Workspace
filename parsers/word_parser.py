"""
Word 文件解析器
"""
from docx import Document  # type: ignore[import-untyped]
from typing import Any, Dict, List, Optional


class WordParser:
    """解析 Word (.docx) 文件"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document: Any = None

    def parse(self) -> Dict[str, Any]:
        """
        解析 Word 文件

        Returns:
            包含標題、段落、表格等結構化內容的字典
        """
        try:
            self.document = Document(self.file_path)

            content = {
                'file_name': self.file_path.split('\\')[-1],
                'file_type': 'docx',
                'paragraphs': self._extract_paragraphs(),
                'tables': self._extract_tables(),
                'headings': self._extract_headings(),
                'full_text': self._extract_full_text()
            }

            return content

        except Exception as e:
            raise Exception(f"Word 文件解析失敗: {str(e)}")

    def _extract_paragraphs(self) -> List[str]:
        """提取所有段落"""
        assert self.document is not None
        paragraphs = []
        for para in self.document.paragraphs:
            text = para.text.strip()
            if text:  # 過濾空段落
                paragraphs.append(text)
        return paragraphs

    def _extract_tables(self) -> List[List[List[str]]]:
        """提取所有表格"""
        assert self.document is not None
        tables = []
        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def _extract_headings(self) -> List[Dict[str, str]]:
        """提取標題結構"""
        assert self.document is not None
        headings = []
        for para in self.document.paragraphs:
            if para.style and para.style.name and para.style.name.startswith('Heading'):
                headings.append({
                    'level': para.style.name,
                    'text': para.text.strip()
                })
        return headings

    def _extract_full_text(self) -> str:
        """提取完整文字（用於 AI 分析）"""
        assert self.document is not None
        full_text = '\n'.join([para.text for para in self.document.paragraphs])
        return full_text.strip()
